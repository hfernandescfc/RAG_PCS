from typing import Tuple, List, Optional

from langchain.docstore.document import Document
from .utils_io import verificar_poppler
from langchain_community.document_loaders import PyPDFLoader
import io
import cv2
import numpy as np

def _to_markdown_table(rows: List[List[Optional[str]]]) -> str:
    # Sanitize and pad rows
    safe = [[(c if c is not None else "").strip() for c in row] for row in rows if row]
    if not safe:
        return ""
    cols = max(len(r) for r in safe)
    norm = [r + ([""] * (cols - len(r))) for r in safe]
    header = norm[0]
    sep = ["---"] * cols
    buf = io.StringIO()
    buf.write("| " + " | ".join(header) + " |\n")
    buf.write("| " + " | ".join(sep) + " |\n")
    for r in norm[1:]:
        buf.write("| " + " | ".join(r) + " |\n")
    return buf.getvalue()


def processar_pdf_com_texto_e_tabelas(pdf_path: str, nome_arquivo: str) -> Tuple[Optional[List[Document]], Optional[str]]:
    """Extrai texto por página e acrescenta tabelas (em Markdown) quando existirem.
    Usa PyPDFLoader para texto e pdfplumber para tabelas.
    """
    try:
        import pdfplumber  # type: ignore
    except Exception:
        # Se pdfplumber não estiver disponível, cai para extração de texto simples
        return processar_pdf_com_texto(pdf_path, nome_arquivo)

    try:
        base_docs = PyPDFLoader(pdf_path).load()
        base_docs = [d for d in base_docs if d.page_content and d.page_content.strip()]

        if not base_docs:
            return None, None

        with pdfplumber.open(pdf_path) as pdf:
            for i, doc in enumerate(base_docs):
                # Ajustar metadados consistentes
                doc.metadata = {
                    "source": nome_arquivo,
                    "page": i + 1,
                    "file_path": pdf_path,
                    "metodo": "extracao_mista",
                }

                if i < len(pdf.pages):
                    page = pdf.pages[i]
                    try:
                        tables = page.extract_tables() or []
                    except Exception:
                        tables = []
                    if tables:
                        md_tables = []
                        for t in tables:
                            md = _to_markdown_table(t)
                            if md:
                                md_tables.append(md)
                        if md_tables:
                            doc.page_content = doc.page_content.strip() + "\n\n[TABELAS DETECTADAS]\n" + "\n\n".join(md_tables)
                            doc.metadata["tem_tabela"] = True
        return base_docs, "extracao_mista"
    except Exception:
        return processar_pdf_com_texto(pdf_path, nome_arquivo)


def verificar_pdf_tem_texto(pdf_path: str) -> bool:
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(pdf_path)
        for page in reader.pages[:3]:
            text = (page.extract_text() or "").strip()
            if len(text) > 100:
                return True
        return False
    except Exception:
        return False


def processar_pdf_com_texto(pdf_path: str, nome_arquivo: str) -> Tuple[Optional[List[Document]], Optional[str]]:
    from langchain_community.document_loaders import PyPDFLoader
    try:
        loader = PyPDFLoader(pdf_path)
        docs = loader.load()
        docs = [d for d in docs if d.page_content and d.page_content.strip()]
        if docs:
            for i, d in enumerate(docs):
                d.metadata = {
                    "source": nome_arquivo,
                    "page": i + 1,
                    "file_path": pdf_path,
                    "metodo": "extracao_direta",
                }
            return docs, "extracao_direta"
    except Exception:
        pass
    return None, None


def processar_pdf_com_ocr(
    pdf_path: str,
    nome_arquivo: str,
    pytesseract,
    convert_from_path,
) -> Tuple[Optional[List[Document]], Optional[str]]:
    import re
    docs: List[Document] = []

    def _ocr_quality(s: str) -> float:
        if not s:
            return 0.0
        chars = len(s)
        alpha = sum(1 for c in s if c.isalpha())
        suspicious = s.count("�") + s.count("□") + s.count("¤")
        score = 0.0
        score += (alpha / chars) if chars else 0.0
        if suspicious > 0:
            score *= 0.8
        return max(0.0, min(1.0, score))

    def _rotate_image_np(img: np.ndarray, angle: int) -> np.ndarray:
        if angle % 360 == 0:
            return img
        if angle == 90:
            return cv2.rotate(img, cv2.ROTATE_90_CLOCKWISE)
        if angle == 180:
            return cv2.rotate(img, cv2.ROTATE_180)
        if angle == 270:
            return cv2.rotate(img, cv2.ROTATE_90_COUNTERCLOCKWISE)
        return img

    def _osd_rotation(pil_img) -> Optional[int]:
        try:
            osd = pytesseract.image_to_osd(pil_img)
            for a in (0, 90, 180, 270):
                if f"Rotate: {a}" in osd:
                    return a
        except Exception:
            return None
        return None

    def _best_rotation(pil_img) -> int:
        best_angle = 0
        best_score = -1
        for ang in (0, 90, 180, 270):
            try:
                test = pil_img.rotate(ang, expand=True)
                txt = pytesseract.image_to_string(test, lang="por", config="--psm 6 -c preserve_interword_spaces=1")
                sc = _ocr_quality(txt)
                if sc > best_score:
                    best_score, best_angle = sc, ang
            except Exception:
                continue
        return best_angle

    def _grid_score(np_img: np.ndarray) -> float:
        # Estima interseções de linhas (grade) normalizadas pela área
        gray = cv2.cvtColor(np_img, cv2.COLOR_BGR2GRAY)
        thr = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]
        h_kernel_len = max(10, np_img.shape[1] // 40)
        v_kernel_len = max(10, np_img.shape[0] // 40)
        h_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (h_kernel_len, 1))
        v_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, v_kernel_len))
        h_lines = cv2.erode(thr, h_kernel, iterations=1)
        h_lines = cv2.dilate(h_lines, h_kernel, iterations=1)
        v_lines = cv2.erode(thr, v_kernel, iterations=1)
        v_lines = cv2.dilate(v_lines, v_kernel, iterations=1)
        inter = cv2.bitwise_and(h_lines, v_lines)
        cnt = int(cv2.countNonZero(inter))
        area = np_img.shape[0] * np_img.shape[1]
        return cnt / max(1, area)

    def _ocr_page_plain(pil_img) -> str:
        return pytesseract.image_to_string(
            pil_img,
            lang="por",
            config="--psm 6 --oem 1 -c preserve_interword_spaces=1",
        )

    def _ocr_table_struct(pil_img) -> Optional[str]:
        # Segmenta grade e realiza OCR célula a célula, produzindo Markdown
        np_img = cv2.cvtColor(np.array(pil_img), cv2.COLOR_RGB2BGR)
        gray = cv2.cvtColor(np_img, cv2.COLOR_BGR2GRAY)
        thr = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]
        h_kernel_len = max(10, np_img.shape[1] // 40)
        v_kernel_len = max(10, np_img.shape[0] // 40)
        h_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (h_kernel_len, 1))
        v_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, v_kernel_len))
        h_lines = cv2.erode(thr, h_kernel, iterations=1)
        h_lines = cv2.dilate(h_lines, h_kernel, iterations=1)
        v_lines = cv2.erode(thr, v_kernel, iterations=1)
        v_lines = cv2.dilate(v_lines, v_kernel, iterations=1)
        inter = cv2.bitwise_and(h_lines, v_lines)

        # Detectar posições únicas de linhas horizontais e verticais via contornos
        def unique_positions(mask: np.ndarray, axis: str) -> List[int]:
            contours, _ = cv2.findContours(mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            vals = []
            for c in contours:
                x, y, w, h = cv2.boundingRect(c)
                if axis == 'h' and w > np_img.shape[1] * 0.3 and h < np_img.shape[0] * 0.05:
                    vals.append(y)
                if axis == 'v' and h > np_img.shape[0] * 0.3 and w < np_img.shape[1] * 0.05:
                    vals.append(x)
            vals.sort()
            # clusterização por proximidade
            uniq = []
            for v in vals:
                if not uniq or abs(v - uniq[-1]) > 5:
                    uniq.append(v)
            return uniq

        ys = unique_positions(h_lines, 'h')
        xs = unique_positions(v_lines, 'v')
        if len(xs) < 2 or len(ys) < 2:
            return None

        rows: List[List[str]] = []
        for r in range(len(ys) - 1):
            row_vals: List[str] = []
            y1, y2 = ys[r], ys[r + 1]
            for c in range(len(xs) - 1):
                x1, x2 = xs[c], xs[c + 1]
                # pequeno padding para evitar linhas
                pad = 2
                y1i, y2i = max(0, y1 + pad), min(np_img.shape[0], y2 - pad)
                x1i, x2i = max(0, x1 + pad), min(np_img.shape[1], x2 - pad)
                if y2i <= y1i or x2i <= x1i:
                    row_vals.append("")
                    continue
                cell = np_img[y1i:y2i, x1i:x2i]
                cell_rgb = cv2.cvtColor(cell, cv2.COLOR_BGR2RGB)
                # PSM 7 ajuda em uma única linha/célula
                txt = pytesseract.image_to_string(cell_rgb, lang="por", config="--psm 7 --oem 1")
                row_vals.append((txt or "").strip())
            rows.append(row_vals)

        md = _to_markdown_table(rows)
        return md if md.strip() else None

    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(pdf_path)
        total_paginas = len(reader.pages)

        poppler = verificar_poppler()
        for page_num in range(1, total_paginas + 1):
            images = convert_from_path(
                pdf_path,
                dpi=400,  # maior DPI para tabelas
                first_page=page_num,
                last_page=page_num,
                fmt="jpeg",
                poppler_path=poppler if poppler else None,
            )
            if not images:
                continue
            pil_img = images[0]

            # 1) Orientação (OSD -> melhor de 0/90/180/270)
            rot = _osd_rotation(pil_img)
            if rot is None:
                rot = _best_rotation(pil_img)
            pil_img = pil_img.rotate(rot, expand=True)

            # 2) Grid score para decidir se é tabela
            np_img = cv2.cvtColor(np.array(pil_img), cv2.COLOR_RGB2BGR)
            gscore = _grid_score(np_img)
            is_table = gscore > 0.005  # limiar empírico

            content = None
            metodo = "ocr"
            tem_tabela = False

            if is_table:
                md = _ocr_table_struct(pil_img)
                if md:
                    content = "[TABELA (OCR estruturado)]\n" + md
                    metodo = "ocr_tabela_struct"
                    tem_tabela = True

            if content is None:
                # OCR normal preservando espaços
                content = (_ocr_page_plain(pil_img) or "").strip()
                # Fallback: se muito ruim e gscore moderado, tenta tabela estruturada
                if _ocr_quality(content) < 0.4 and gscore > 0.002:
                    md2 = _ocr_table_struct(pil_img)
                    if md2:
                        content = "[TABELA (OCR estruturado)]\n" + md2
                        metodo = "ocr_tabela_struct"
                        tem_tabela = True

            if content and content.strip():
                docs.append(
                    Document(
                        page_content=content.strip(),
                        metadata={
                            "source": nome_arquivo,
                            "page": page_num,
                            "file_path": pdf_path,
                            "metodo": metodo,
                            "tem_tabela": tem_tabela,
                            "rotacao": rot,
                            "grid_score": float(gscore),
                        },
                    )
                )
    except Exception:
        return None, None
    return (docs, "ocr") if docs else (None, None)


def processar_docx(docx_path: str, nome_arquivo: str) -> Tuple[Optional[List[Document]], Optional[str]]:
    try:
        from langchain_community.document_loaders import Docx2txtLoader
        loader = Docx2txtLoader(docx_path)
        documents = loader.load()
        documents = [doc for doc in documents if doc.page_content and doc.page_content.strip()]
        if documents:
            for doc in documents:
                doc.metadata = {
                    "source": nome_arquivo,
                    "page": 1,
                    "file_path": docx_path,
                    "metodo": "docx",
                }
            return documents, "docx"
    except Exception:
        pass

    try:
        import docx as docx_lib
        d = docx_lib.Document(docx_path)
        text = "\n".join(p.text for p in d.paragraphs if p.text and p.text.strip())
        if text.strip():
            return [Document(page_content=text.strip(), metadata={
                "source": nome_arquivo,
                "page": 1,
                "file_path": docx_path,
                "metodo": "docx",
            })], "docx"
    except Exception:
        pass

    return None, None
