import os
import sys
import glob
import hashlib
import json
import gc
from rag_pipeline.config import get_embeddings, DB_PATH
from rag_pipeline import utils_io as _io
from rag_pipeline import utils_pdf as _pdf

def verificar_tesseract():
    caminhos_possiveis = [
        r'C:\Program Files\Tesseract-OCR\tesseract.exe',
        r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
        r'C:\Tesseract-OCR\tesseract.exe',
    ]
    for caminho in caminhos_possiveis:
        if os.path.exists(caminho):
            return caminho
    return None

def verificar_poppler():
    """Retorna caminho do binário do Poppler (Windows) se encontrado.
    Usa a env var POPPLER_PATH quando definida e alguns caminhos comuns.
    """
    # Preferir variável de ambiente
    env_path = os.getenv('POPPLER_PATH')
    if env_path and os.path.isdir(env_path):
        return env_path

    # Caminhos comuns no Windows
    candidatos = [
        r'C:\\Program Files\\poppler\\Library\\bin',
        r'C:\\Program Files\\poppler-24.08.0\\Library\\bin',
        r'C:\\Program Files\\poppler-24.07.0\\Library\\bin',
        r'C:\\Program Files\\poppler-24.02.0\\Library\\bin',
        r'C:\\poppler\\Library\\bin',
        r'C:\\poppler\\bin',
    ]
    for caminho in candidatos:
        if os.path.isdir(caminho):
            return caminho
    return None

# Caminho global do Poppler (usado no OCR com pdf2image)
POPPLER_PATH = verificar_poppler()

def calcular_hash_arquivo(filepath):
    hash_md5 = hashlib.md5()
    with open(filepath, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_md5.update(chunk)
    return hash_md5.hexdigest()

def carregar_registro_processados(registro_path="./arquivos_processados.json"):
    if os.path.exists(registro_path):
        with open(registro_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

def salvar_registro_processados(registro, registro_path="./arquivos_processados.json"):
    with open(registro_path, 'w', encoding='utf-8') as f:
        json.dump(registro, f, indent=2, ensure_ascii=False)

def verificar_pdf_tem_texto(pdf_path):
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(pdf_path)
        for i, page in enumerate(reader.pages[:3]):
            text = page.extract_text().strip()
            if len(text) > 100:
                return True
        return False
    except:
        return False

def processar_pdf_com_texto(pdf_path, nome_arquivo):
    from langchain_community.document_loaders import PyPDFLoader
    from langchain.docstore.document import Document
    
    print("   📝 Extração direta...")
    
    try:
        loader = PyPDFLoader(pdf_path)
        documents = loader.load()
        documents = [doc for doc in documents if doc.page_content.strip()]
        
        if documents:
            # CORREÇÃO: Garantir metadados corretos
            for i, doc in enumerate(documents):
                doc.metadata = {
                    'source': nome_arquivo,  # Nome do arquivo (não caminho!)
                    'page': i + 1,
                    'file_path': pdf_path,
                    'metodo': 'extração_direta'
                }
            return documents, 'extração_direta'
    except Exception as e:
        print(f"   ⚠️  Erro: {e}")
    
    return None, None

def processar_docx(docx_path, nome_arquivo):
    """Extrai texto de arquivos Word (.docx) e retorna lista de Documents."""
    from langchain.docstore.document import Document
    print("   -> Processando DOCX...")
    # Tentar com loader do LangChain (docx2txt)
    try:
        from langchain_community.document_loaders import Docx2txtLoader
        loader = Docx2txtLoader(docx_path)
        documents = loader.load()
        documents = [doc for doc in documents if doc.page_content and doc.page_content.strip()]
        if documents:
            for doc in documents:
                doc.metadata = {
                    'source': nome_arquivo,
                    'page': 1,
                    'file_path': docx_path,
                    'metodo': 'docx'
                }
            return documents, 'docx'
    except Exception:
        pass

    # Fallback com python-docx
    try:
        import docx as docx_lib
        d = docx_lib.Document(docx_path)
        text = "\n".join(p.text for p in d.paragraphs if p.text and p.text.strip())
        if text.strip():
            return [Document(page_content=text.strip(), metadata={
                'source': nome_arquivo,
                'page': 1,
                'file_path': docx_path,
                'metodo': 'docx'
            })], 'docx'
    except Exception as e:
        print(f"   !! Erro DOCX: {e}")
    return None, None

def processar_pdf_com_ocr(pdf_path, nome_arquivo, pytesseract, convert_from_path):
    from langchain.docstore.document import Document
    
    print("   🖼️  OCR (página por página)...")
    
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(pdf_path)
        total_paginas = len(reader.pages)
        print(f"   Total: {total_paginas} páginas")
        
        documents = []
        
        for page_num in range(1, total_paginas + 1):
            try:
                percent = (page_num / total_paginas) * 100
                print(f"      [{page_num}/{total_paginas}] {percent:.0f}%", end="\r")
                
                # Usar Poppler quando disponível (necessário no Windows)
                if POPPLER_PATH:
                    images = convert_from_path(
                        pdf_path,
                        dpi=200,
                        first_page=page_num,
                        last_page=page_num,
                        fmt='jpeg',
                        poppler_path=POPPLER_PATH
                    )
                else:
                    images = convert_from_path(
                        pdf_path,
                        dpi=200,
                        first_page=page_num,
                        last_page=page_num,
                        fmt='jpeg'
                    )
                
                if images:
                    image = images[0]
                    text = pytesseract.image_to_string(
                        image,
                        lang='por',
                        config='--psm 3 --oem 1'
                    )
                    
                    if text.strip():
                        # CORREÇÃO: Metadados explícitos e corretos
                        doc = Document(
                            page_content=text.strip(),
                            metadata={
                                "source": nome_arquivo,  # USAR NOME DO ARQUIVO
                                "page": page_num,
                                "file_path": pdf_path,
                                "metodo": "ocr"
                            }
                        )
                        documents.append(doc)
                    
                    del images
                    del image
                    gc.collect()
                    
            except Exception as e:
                print(f"\n      ⚠️  Erro página {page_num}: {e}")
        
        print(f"\n   ✅ {len(documents)} páginas extraídas")
        return documents, 'ocr'
        
    except Exception as e:
        print(f"   ❌ Erro no OCR: {e}")
        return None, None

# ---------------------------------------------------------------------------
# Overrides para usar utilitários centralizados (evita divergência de código)
# ---------------------------------------------------------------------------
def _override_verificar_tesseract():
    return _io.verificar_tesseract()

def _override_verificar_poppler():
    return _io.verificar_poppler()

def _override_calcular_hash_arquivo(filepath):
    return _io.calcular_hash_arquivo(filepath)

def _override_carregar_registro_processados(registro_path="./arquivos_processados.json"):
    return _io.carregar_registro_processados(registro_path)

def _override_salvar_registro_processados(registro, registro_path="./arquivos_processados.json"):
    return _io.salvar_registro_processados(registro, registro_path)

def _override_verificar_pdf_tem_texto(pdf_path: str) -> bool:
    return _pdf.verificar_pdf_tem_texto(pdf_path)

def _override_processar_pdf_com_texto(pdf_path: str, nome_arquivo: str):
    return _pdf.processar_pdf_com_texto(pdf_path, nome_arquivo)

def _override_processar_docx(docx_path: str, nome_arquivo: str):
    return _pdf.processar_docx(docx_path, nome_arquivo)

def _override_processar_pdf_com_ocr(pdf_path, nome_arquivo, pytesseract, convert_from_path):
    return _pdf.processar_pdf_com_ocr(pdf_path, nome_arquivo, pytesseract, convert_from_path)

# Reatribuir nomes originais para usar as implementações centralizadas
verificar_tesseract = _override_verificar_tesseract
verificar_poppler = _override_verificar_poppler
calcular_hash_arquivo = _override_calcular_hash_arquivo
carregar_registro_processados = _override_carregar_registro_processados
salvar_registro_processados = _override_salvar_registro_processados
verificar_pdf_tem_texto = _override_verificar_pdf_tem_texto
processar_pdf_com_texto = _override_processar_pdf_com_texto
processar_docx = _override_processar_docx
processar_pdf_com_ocr = _override_processar_pdf_com_ocr

def processar_e_salvar_arquivo(arq_info, registro, embeddings, db_path, ocr_disponivel, 
                               pytesseract=None, convert_from_path=None):
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain_community.vectorstores import Chroma
    from datetime import datetime
    
    pdf_path = arq_info['path']
    nome_arquivo = arq_info['nome']
    hash_arquivo = arq_info['hash']
    
    print(f"\n{'='*70}")
    print(f"📄 {nome_arquivo}")
    print('='*70)
    
    # Verificar se já foi processado
    if hash_arquivo in registro:
        print("   ⏭️  Já processado - pulando")
        return True
    
    # Roteamento por extensão
    ext = os.path.splitext(pdf_path)[1].lower()

    documents = None
    metodo = None

    if ext == '.pdf':
        print("1??  Verificando texto...")
        tem_texto = verificar_pdf_tem_texto(pdf_path)
        if tem_texto:
            print("   ? Tem texto!")
            documents, metodo = processar_pdf_com_texto(pdf_path, nome_arquivo)
        if not documents and ocr_disponivel:
            print("2??  Usando OCR...")
            documents, metodo = processar_pdf_com_ocr(pdf_path, nome_arquivo, 
                                                      pytesseract, convert_from_path)
    elif ext == '.docx':
        documents, metodo = processar_docx(pdf_path, nome_arquivo)
    else:
        print(f"   !! Extensão não suportada: {ext}")
    if not documents:
        print("   ❌ Falhou")
        return False
    
    # VERIFICAÇÃO: Conferir metadados antes de salvar
    print(f"   🔍 Verificando metadados...")
    sources_unicos = set(doc.metadata.get('source') for doc in documents)
    print(f"   ✅ Source único: {sources_unicos}")
    
    if len(sources_unicos) != 1 or nome_arquivo not in sources_unicos:
        print(f"   ⚠️  AVISO: Source inconsistente!")
        print(f"      Esperado: {nome_arquivo}")
        print(f"      Encontrado: {sources_unicos}")
    
    # Dividir em chunks
    print("   ✂️  Criando chunks...")
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,  # AUMENTADO de 500 para 1000
        chunk_overlap=100  # AUMENTADO de 50 para 100
    )
    docs = text_splitter.split_documents(documents)
    docs = [doc for doc in docs if doc.page_content.strip()]
    
    # VERIFICAÇÃO: Conferir metadados dos chunks
    sources_chunks = set(doc.metadata.get('source') for doc in docs)
    print(f"   ✅ {len(docs)} chunks criados")
    print(f"   ✅ Sources dos chunks: {sources_chunks}")
    
    # Salvar no banco
    print("   💾 Salvando...")
    
    try:
        if os.path.exists(db_path):
            db = Chroma(persist_directory=db_path, embedding_function=embeddings)
            db.add_documents(docs)
        else:
            db = Chroma.from_documents(docs, embeddings, persist_directory=db_path)
        
        print("   ✅ Salvo!")
        
        # Registrar sucesso
        registro[hash_arquivo] = {
            'nome': nome_arquivo,
            'caminho': pdf_path,
            'metodo': metodo,
            'paginas': len(documents),
            'chunks': len(docs),
            'data_processo': datetime.now().isoformat()
        }
        salvar_registro_processados(registro)
        
        del documents
        del docs
        del db
        gc.collect()
        
        return True
        
    except Exception as e:
        print(f"   ❌ Erro ao salvar: {e}")
        return False

print("🔧 Sistema RAG - VERSÃO CORRIGIDA (Metadados Preservados)\n")

# Verificar Tesseract
tesseract_path = verificar_tesseract()

if not tesseract_path:
    print("⚠️  Tesseract não encontrado - OCR desabilitado")
    ocr_disponivel = False
else:
    print(f"✅ Tesseract: {tesseract_path}")
    ocr_disponivel = True

# Importar bibliotecas
try:
    from langchain_community.document_loaders import PyPDFLoader
    from PyPDF2 import PdfReader
    import requests
    from langchain_community.vectorstores import Chroma
    from langchain_community.embeddings import HuggingFaceEmbeddings
    from langchain_community.llms import Ollama
    from langchain.chains import RetrievalQA
except ImportError as e:
    print(f"❌ Erro: {e}")
    sys.exit(1)

if ocr_disponivel:
    try:
        from pdf2image import convert_from_path
        import pytesseract
        pytesseract.pytesseract.tesseract_cmd = tesseract_path
    except ImportError:
        print("⚠️  pdf2image/pytesseract não disponíveis")
        ocr_disponivel = False
        pytesseract = None
        convert_from_path = None

print()

# ============================================================================
# CONFIGURAÇÃO
# ============================================================================

# OPÇÃO 1: Lista manual
# pdf_files = [
#     r"C:\Users\compesa\Desktop\ppp_RAG\doc1.pdf",
# ]

# OPÇÃO 2: Pasta inteira (RECOMENDADO)
pdf_folder = r"C:\Users\compesa\Desktop\ppp_RAG"

# A busca agora é recursiva para encontrar arquivos .pdf e .PDF em todas as subpastas
# O '**' indica que a busca deve incluir todos os subdiretórios
pdf_folder = os.path.join(os.getcwd(), "data", "raw")  # override para pasta de dados
pdf_files = glob.glob(os.path.join(pdf_folder, "**", "*.pdf"), recursive=True)
pdf_files += glob.glob(os.path.join(pdf_folder, "**", "*.PDF"), recursive=True)
# Incluir arquivos Word (.docx)
pdf_files += glob.glob(os.path.join(pdf_folder, "**", "*.docx"), recursive=True)
pdf_files += glob.glob(os.path.join(pdf_folder, "**", "*.DOCX"), recursive=True)
# ============================================================================

pdf_files = [f for f in pdf_files if os.path.exists(f)]

# Carregar registro
registro = carregar_registro_processados()
print(f"📋 Registro: {len(registro)} arquivos já processados\n")

# Verificar o que processar
arquivos_para_processar = []
hashes_vistos = set()

print("🔍 Verificando arquivos...\n")

for pdf_path in pdf_files:
    nome_arquivo = os.path.basename(pdf_path)
    hash_arquivo = calcular_hash_arquivo(pdf_path)
    
    if hash_arquivo in registro:
        print(f"⏭️  {nome_arquivo}")
        continue
    
    if hash_arquivo in hashes_vistos:
        print(f"⚠️  {nome_arquivo} (duplicado)")
        continue
    
    hashes_vistos.add(hash_arquivo)
    size_mb = os.path.getsize(pdf_path) / (1024*1024)
    print(f"✅ {nome_arquivo} ({size_mb:.2f} MB)")
    
    arquivos_para_processar.append({
        'path': pdf_path,
        'nome': nome_arquivo,
        'hash': hash_arquivo
    })

print(f"\n{'='*70}")
print(f"📊 Para processar: {len(arquivos_para_processar)} arquivos")
print('='*70)

if not arquivos_para_processar:
    print("\n✅ Todos já processados!")
    pass

# Criar embeddings
print("\n🧠 Inicializando embeddings...")
embeddings = get_embeddings()
print("✅ Embeddings prontos")

db_path = DB_PATH

# Conciliação: adicionar para reprocessar itens do registro que estão ausentes no banco
try:
    from langchain_community.vectorstores import Chroma
    missing_added = 0
    existing_sources = set()
    if os.path.exists(db_path):
        db_chk = Chroma(persist_directory=db_path, embedding_function=embeddings)
        data_chk = db_chk.get()
        existing_sources = set(m.get('source') for m in data_chk['metadatas'])
        del db_chk
    # Hashes já na fila
    fila_hashes = {item['hash'] for item in arquivos_para_processar}
    for h, info in registro.items():
        nome = info.get('nome')
        caminho = info.get('caminho')
        if nome and (nome not in existing_sources) and caminho and os.path.exists(caminho):
            if h not in fila_hashes:
                arquivos_para_processar.append({
                    'path': caminho,
                    'nome': os.path.basename(caminho),
                    'hash': h
                })
                missing_added += 1
    if missing_added:
        print(f"Adicionados {missing_added} arquivos do registro para reindexação (ausentes no banco)")
except Exception as _:
    pass

# Processar arquivos
sucesso = 0
falhas = 0

for i, arq_info in enumerate(arquivos_para_processar, 1):
    print(f"\n[{i}/{len(arquivos_para_processar)}]")
    
    try:
        if processar_e_salvar_arquivo(
            arq_info, registro, embeddings, db_path, 
            ocr_disponivel, pytesseract, convert_from_path
        ):
            sucesso += 1
        else:
            falhas += 1
    except Exception as e:
        print(f"   ❌ ERRO: {e}")
        falhas += 1
    
    gc.collect()

print(f"\n{'='*70}")
print(f"✅ PROCESSAMENTO FINALIZADO")
print('='*70)
print(f"✅ Sucesso: {sucesso}")
print(f"❌ Falhas: {falhas}")

# VERIFICAÇÃO FINAL
print(f"\n{'='*70}")
print("🔍 VERIFICAÇÃO FINAL DO BANCO")
print('='*70)

if os.path.exists(db_path):
    print("\n💾 Carregando banco para verificação...")
    db = Chroma(persist_directory=db_path, embedding_function=embeddings)
    all_data = db.get()
    
    from collections import Counter
    sources = [meta.get('source', 'SEM_SOURCE') for meta in all_data['metadatas']]
    source_counts = Counter(sources)
    
    print(f"\n📊 RESULTADO:")
    print(f"   Total de chunks: {len(all_data['ids'])}")
    print(f"   Fontes únicas: {len(source_counts)}")
    
    if len(source_counts) > 1:
        print(f"\n✅ SUCESSO! Múltiplas fontes detectadas:")
        for source, count in sorted(source_counts.items())[:20]:
            print(f"   • {source:50s} → {count:4d} chunks")
    else:
        print(f"\n❌ PROBLEMA PERSISTE!")
        only = list(source_counts.keys())[0] if len(source_counts) == 1 else "NENHUMA"
        print(f"   Apenas 1 fonte: {only}")
    
    print(f"\n{'='*70}")

print("\n✅ Script concluído!")


